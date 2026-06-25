from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/choi/badamoyeo")
OUT = ROOT / "docs" / "badamoyeo_ai_usage_report.docx"

BODY_FONT = "Malgun Gothic"
CODE_FONT = "Consolas"
BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F6F7F9"
BORDER = "D9E2EC"


def set_run_font(run, name=BODY_FONT, size=10.5, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn("w:" + margin))
        if element is None:
            element = OxmlElement("w:" + margin)
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 10, 5),
        ("Heading 3", 11.5, RGBColor(38, 38, 38), 8, 3),
    ]:
        style = doc.styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("바다모여 AI 활용 보고서")
    set_run_font(run, size=8.5, color=MUTED)


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("바다모여 AI 활용 보고서")
    set_run_font(run, size=23, color=RGBColor(0, 0, 0), bold=True)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(12)
    run = p2.add_run("GPT 기반 챗봇, 메인 추천, 스팟 상세 분석 구현 정리")
    set_run_font(run, size=11.5, color=MUTED)

    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, [1600, 7600])
    rows = [
        ("대상", "바다모여 백엔드 애플리케이션"),
        ("AI 적용 지점", "AI 챗봇 / 메인 화면 AI 추천 / 스팟 상세 AI 분석"),
        ("공통 모델", "gpt-5-mini, GMS Responses API"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        for cell in row.cells:
            set_cell_borders(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(label), size=9.5, bold=True)
        p = row.cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(value), size=9.5)


def add_summary(doc):
    doc.add_heading("1. AI 활용 개요", level=1)
    p = doc.add_paragraph()
    set_run_font(
        p.add_run(
            "본 애플리케이션은 해양 레저 장소 추천과 예보 해석 품질을 높이기 위해 AI를 세 곳에 적용하였다. "
            "AI는 사용자의 자연어 질문을 이해하고, 해양 예보 DB 후보를 평가하며, 특정 장소의 예보를 짧은 추천 사유로 변환한다."
        )
    )

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, [1550, 2450, 2750, 2650])
    headers = ["구분", "AI 역할", "입력 데이터", "결과"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_borders(cell)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(text), size=9, bold=True)
    rows = [
        ("AI 챗봇", "질문 의도 판단 및 장소 검색 도구 호출", "사용자 질문, 서버 기준 날짜", "DB 예보 기반 짧은 답변"),
        ("메인 AI 추천", "후보 장소 비교 및 추천 순위 생성", "체험별 장소 후보와 오전/오후/일 예보", "추천 순위와 이유 저장"),
        ("상세 AI 분석", "특정 예보의 방문 추천 여부 판단", "장소, 날짜, 시간대, 종합지수, 지표", "추천 여부와 1문장 이유"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            set_cell_borders(cells[idx])
            set_cell_margins(cells[idx])
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(value), size=8.7)


def add_code_box(doc, title, path, lines, code):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{title}  |  {path}:{lines}")
    set_run_font(run, size=9.2, color=BLUE, bold=True)

    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_borders(cell, "CBD5E1")
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(code.strip("\n").split("\n")):
        if idx:
            p.add_run("\n")
        run = p.add_run(line)
        set_run_font(run, name=CODE_FONT, size=8.0, color=RGBColor(31, 41, 55))


def add_note(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF8E6")
    set_cell_borders(cell, "E7C873")
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(label + ": "), size=9.5, bold=True, color=RGBColor(122, 90, 0))
    set_run_font(p.add_run(text), size=9.5)


def add_feature(doc, num, title, overview, prompt_code, core_codes, flow, effect):
    doc.add_heading(f"{num}. {title}", level=1)
    p = doc.add_paragraph()
    set_run_font(p.add_run(overview))
    doc.add_heading("프롬프트", level=2)
    add_code_box(doc, "프롬프트 캡처", prompt_code["path"], prompt_code["lines"], prompt_code["code"])
    doc.add_heading("주요 코드", level=2)
    for code in core_codes:
        add_code_box(doc, code["title"], code["path"], code["lines"], code["code"])
    doc.add_heading("동작 설명", level=2)
    for sentence in flow:
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["List Bullet"]
        set_run_font(p.add_run(sentence))
    add_note(doc, "활용 효과", effect)


def add_conclusion(doc):
    doc.add_heading("5. AI 활용 효과 및 정리", level=1)
    points = [
        "사용자는 복잡한 필터를 직접 선택하지 않아도 자연어로 장소 추천을 받을 수 있다.",
        "메인 추천은 AI가 여러 시간대 예보와 게시글 수를 함께 고려해 추천 순위와 이유를 생성한다.",
        "상세 분석은 특정 장소의 예보를 사용자가 이해하기 쉬운 추천 여부와 한 문장 이유로 변환한다.",
        "추천 및 분석 결과는 DB에 저장하거나 최신 여부를 확인해 재사용하므로 불필요한 AI 호출을 줄인다.",
        "프롬프트에서 DB에 없는 장소, 수치, 시설 정보는 추측하지 않도록 제한하여 답변 신뢰도를 높였다.",
    ]
    for point in points:
        p = doc.add_paragraph(style="List Bullet")
        set_run_font(p.add_run(point))


def main():
    doc = Document()
    style_doc(doc)
    add_title(doc)
    add_summary(doc)

    add_feature(
        doc,
        2,
        "AI 챗봇 기능",
        "사용자의 자연어 질문을 GPT 모델에 전달하고, 장소 추천이나 예보 확인이 필요한 경우 AI가 DB 검색 도구를 호출하도록 구현하였다.",
        {
            "path": "GmsAiChatbotService.java",
            "lines": "50-98",
            "code": '''
private static final String SYSTEM_PROMPT = """
    당신은 대한민국 해양 레저 앱 '바다모여'의 안내 챗봇입니다.
    바다 여행, 해수욕, 갯벌 체험, 스쿠버다이빙, 낚시, 서핑에 관해
    정확하고 다정한 한국어로 답하세요.

    [모바일 답변 원칙]
    - 결론부터 말하고 전체 답변은 500자 이내, 최대 7문장으로 작성하세요.
    - 한 문장은 짧게 쓰고, 인사말·질문 반복·불필요한 서론은 생략하세요.
    - 장소를 나열할 때만 줄바꿈과 '- ' 목록을 사용하고 표와 긴 제목은 쓰지 마세요.
    - 이모지는 사용하지 마세요.
    - 사용자가 요청하지 않은 정보나 후속 제안은 덧붙이지 마세요.
    - 함수 호출 과정, 함수명, 인자, JSON, DB 원문, 영문 필드명은 절대 답변에 노출하지 마세요.
    - '조회하겠습니다' 같은 처리 과정은 말하지 말고 조회가 끝난 최종 답변만 작성하세요.

    [질문 처리]
    1. 일반적인 준비물, 체험 방법, 안전수칙 질문은 알고 있는 범위에서 바로 답하세요.
    2. 답변에 꼭 필요한 지역이나 체험 종류가 없으면 추측하지 말고 한 번에 한 가지만 짧게 물으세요.
    3. 장소 추천·비추천, 특정 지역·장소의 상태나 예보를 묻는 질문은
       반드시 searchMarineSpots로 바다모여 DB를 먼저 조회하세요.
    4. 좋은 곳과 피할 곳을 함께 요청하면 best와 worst를 각각 조회하세요.
       추천만 요청하면 best만, 비추천만 요청하면 worst만 조회하세요.
    5. 사용자가 지정한 날짜는 targetDate에 YYYY-MM-DD로 전달하세요.
       오늘·내일·주말은 입력에 제공된 서버 기준 오늘 날짜로 계산하고,
       날짜를 지정하지 않은 경우에만 빈 문자열을 전달하세요.

    [DB 답변 규칙]
    - 함수가 반환한 장소명, 날짜, 시간대, 종합지수, 날씨, 물때, 세부 지표만 사용하세요.
    - DB에 없는 장소, 수치, 시설, 교통, 운영시간을 추측하거나 만들지 마세요.
    - 결과가 없으면 '조건에 맞는 예보 데이터를 찾지 못했어요.'라고 말하세요.
      지역 또는 체험 종류가 빠진 경우에만 필요한 정보 하나를 짧게 물으세요.
    - 기준 날짜는 결론에 한 번만 밝히고, 장소는 중요도순 최대 2곳만 안내하세요.
    - 각 장소는 '장소명(시간대) — 종합지수. 판단 이유' 형식의 한 줄로 쓰세요.
    - 판단 이유에는 해당 체험과 관련 있는 실제 지표를 가능하면 2개 사용하고,
      수치만 나열하지 말고 그 조건이 체험에 왜 유리하거나 불리한지 친절히 설명하세요.
    - 지표의 의미를 확실히 판단할 수 없으면 과장하지 말고 관찰된 수치만 안내하세요.
    - '매우좋음'과 '좋음'만 추천으로 표현하세요.
    - '보통'은 조건부 방문 가능으로만 표현하세요.
    - '나쁨'과 '매우나쁨'은 추천하지 말고 피해야 할 이유로만 설명하세요.
    - best는 상대 정렬일 뿐이므로 종합지수를 확인한 뒤 추천 여부를 판단하세요.
    - 결과가 모두 '나쁨' 또는 '매우나쁨'이면
      '해당 날짜에는 추천할 장소가 없습니다.'로 시작하세요.
    - 종합지수가 '나쁨' 또는 '매우나쁨'인 장소에는 '추천', '적합',
      '도전 가능', '수련 가능' 같은 긍정 표현을 어떤 경우에도 사용하지 마세요.

    [안전]
    - 안전을 보장하거나 실시간 상태를 알고 있는 것처럼 말하지 마세요.
    - 장소·예보 답변의 마지막에는 '방문 전 최신 예보와 현장 통제를 확인하세요.'를 한 번만 쓰세요.
    - 안전 질문에는 구명조끼 착용과 현장 안전수칙 준수를 우선 안내하세요.
    """;
''',
        },
        [
            {
                "title": "사용자 질문을 AI 요청으로 변환",
                "path": "GmsAiChatbotService.java",
                "lines": "117-127, 161-168",
                "code": '''
public ChatCompletionResponse complete(String message) {
    JsonNode firstResponse = responsesClient.execute(initialRequest(message.trim()));
    List<ToolSearchResult> toolOutputs = executeToolCalls(firstResponse);

    String content = toolOutputs.isEmpty()
        ? responsesClient.outputText(firstResponse)
        : formatToolAnswer(toolOutputs);

    return new ChatCompletionResponse(normalizeForChatBubble(content), Instant.now());
}

private Map<String, Object> initialRequest(String message) {
    String datedInput = """
        서버 기준 오늘 날짜: %s
        사용자 질문: %s
        """.formatted(LocalDate.now(), message);
    Map<String, Object> request = responsesClient.baseRequest(model, SYSTEM_PROMPT, datedInput, 700);
    request.put("tools", List.of(searchMarineSpotsTool()));
    return request;
}
''',
            },
            {
                "title": "AI function call 결과로 DB 검색 실행",
                "path": "GmsAiChatbotService.java",
                "lines": "171-189",
                "code": '''
JsonNode arguments = objectMapper.readTree(output.path("arguments").asText("{}"));
List<AiSpotSearchResult> result = spotRecommendationTools.searchMarineSpots(
    arguments.path("experience").asText(""),
    arguments.path("region").asText(""),
    arguments.path("keyword").asText(""),
    arguments.path("targetDate").asText(""),
    arguments.path("sort").asText("best"),
    Math.min(arguments.path("limit").asInt(2), 2)
);
''',
            },
        ],
        [
            "챗봇은 사용자 질문과 서버 기준 날짜를 함께 AI에 전달한다.",
            "장소 추천이 필요한 질문이면 AI가 `searchMarineSpots` 도구 호출을 생성한다.",
            "서버는 AI가 만든 조건을 파싱한 뒤 실제 DB를 조회하고, 조회 결과를 짧은 모바일 답변으로 정리한다.",
        ],
        "사용자가 자연어로 질문해도 DB 예보를 기반으로 추천 장소와 이유를 받을 수 있으며, AI가 DB에 없는 정보를 추측하지 않도록 제한하였다.",
    )

    add_feature(
        doc,
        3,
        "메인 화면 AI 추천 기능",
        "해양 예보 수집 이후 체험별 후보 장소를 AI에게 전달하고, 추천 순위와 이유를 생성해 DB에 저장한다.",
        {
            "path": "GmsAiRecommendationGenerator.java",
            "lines": "20-30",
            "code": '''
private static final String SYSTEM_PROMPT = """
    당신은 바다모여의 해양 레저 장소 추천 분석가입니다.
    제공된 DB 후보만 평가하고 새로운 장소나 수치를 만들지 마세요.
    각 장소의 해당 날짜 오전·오후·일 예보 전체를 종합해서 평가하세요.
    특정 시간대 하나만 보고 판단하지 말고 시간대별 종합지수, 날씨, 물때,
    체험별 세부 지표와 게시글 수를 함께 고려하세요.
    반드시 {"recommendations":[{"spotId":1,"rank":1,"reason":"이유"}]} 형태의 JSON만 반환하세요.
    """;
''',
        },
        [
            {
                "title": "후보 데이터를 AI에 전달하고 JSON 추천 결과 수신",
                "path": "GmsAiRecommendationGenerator.java",
                "lines": "47-63",
                "code": '''
AiRecommendationSelection selection = responsesClient.completeJson(
    model,
    SYSTEM_PROMPT,
    """
        체험 종류: %s
        아래 후보 중 추천 장소를 정확히 %d곳 선정하세요.
        rank는 1부터 시작해 중복 없이 부여하고, reason은 100자 이내로 작성하세요.

        후보 데이터:
        %s
        """.formatted(experience, count, objectMapper.writeValueAsString(candidates)),
    1200,
    AiRecommendationSelection.class
);
''',
            },
            {
                "title": "AI 추천 결과 검증 후 DB 저장",
                "path": "AiRecommendationService.java",
                "lines": "65-98",
                "code": '''
List<AiRecommendationCandidate> candidates =
    recommendationMapper.findCandidates(experience, recommendationDate, CANDIDATE_LIMIT);

List<AiRecommendationItem> generated =
    recommendationGenerator.generate(experience, candidates);

List<AiRecommendationSaveRequest> recommendations =
    validateAndConvert(experience, recommendationDate, candidates, generated);

transactionTemplate.executeWithoutResult(status -> {
    recommendationMapper.deleteRecommendations(experience, recommendationDate);
    recommendationMapper.insertRecommendations(recommendations);
});
''',
            },
        ],
        [
            "DB에서 체험과 날짜에 맞는 후보 장소를 최대 20개 조회한다.",
            "AI는 후보 중 추천 장소 6곳을 선정하고, 각 장소의 순위와 추천 이유를 JSON으로 반환한다.",
            "서버는 AI 결과의 장소 ID, 순위 중복, 이유 길이를 검증한 뒤 추천 테이블에 저장한다.",
        ],
        "메인 화면 조회 때마다 AI를 호출하지 않고 저장된 추천 결과를 재사용하여 응답 속도와 AI 사용 비용을 줄였다.",
    )

    add_feature(
        doc,
        4,
        "스팟 상세 AI 분석 기능",
        "특정 장소의 예보 데이터를 AI가 분석하여 방문 추천 여부와 한 문장 이유를 생성한다.",
        {
            "path": "GmsAiSpotAnalysisGenerator.java",
            "lines": "17-25",
            "code": '''
private static final String SYSTEM_PROMPT = """
    당신은 바다모여의 해양 레저 장소 분석가입니다.
    제공된 DB 데이터만 사용하고 시설, 교통, 운영시간 또는 수치를 추측하지 마세요.
    종합지수와 체험별 지표를 함께 분석하되 안전을 보장한다고 표현하지 마세요.
    recommended는 종합지수가 '매우좋음' 또는 '좋음'일 때만 true로 판단하세요.
    recommendationReason은 실제 지표를 1~2개 사용해 한국어 1문장, 120자 이내로 작성하세요.
    반드시 {"recommended":true,"recommendationReason":"이유"} 형태의 JSON만 반환하세요.
    """;
''',
        },
        [
            {
                "title": "특정 예보 데이터를 AI에 전달",
                "path": "GmsAiSpotAnalysisGenerator.java",
                "lines": "42-55",
                "code": '''
return responsesClient.completeJson(
    model,
    SYSTEM_PROMPT,
    """
        다음 장소의 해당 예보를 분석하세요.
        recommended는 이 체험을 위해 방문을 추천할 수 있는 데이터인지 판단한 값입니다.

        %s
        """.formatted(objectMapper.writeValueAsString(source)),
    800,
    AiSpotAnalysisContent.class
);
''',
            },
            {
                "title": "기존 분석 재사용 또는 새 분석 저장",
                "path": "AiSpotAnalysisService.java",
                "lines": "59-75",
                "code": '''
AiSpotAnalysisRow cached = analysisMapper.findAnalysis(source.forecastId());
if (cached != null && source.forecastUpdatedAt().equals(cached.sourceForecastUpdatedAt())) {
    return toResponse(cached);
}

AiSpotAnalysisContent generated = validate(analysisGenerator.generate(source));
analysisMapper.upsertAnalysis(new AiSpotAnalysisSaveRequest(
    source.spotId(), source.forecastId(), generated.recommended(),
    generated.recommendationReason().trim(), source.forecastUpdatedAt(), generatedAt
));
''',
            },
        ],
        [
            "상세 분석 API는 장소와 날짜에 맞는 예보 원본 데이터를 조회한다.",
            "기존 분석이 있고 원본 예보의 수정 시간이 같으면 저장된 결과를 재사용한다.",
            "새 분석이 필요할 때만 AI를 호출하고, 추천 여부와 이유를 `ai_spot_analyses`에 저장한다.",
        ],
        "사용자는 단순한 종합지수뿐 아니라 해당 장소를 방문해도 되는지와 그 이유를 한 문장으로 확인할 수 있다.",
    )

    add_conclusion(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    main()
